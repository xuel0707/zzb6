if False:
    import os
    import json
    from docx import Document
    import subprocess

    def extract_text_from_docx(docx_path):
        """Extract text from a .docx file and return as a single string."""
        document = Document(docx_path)
        full_text = []

        # Extract paragraphs
        for para in document.paragraphs:
            full_text.append(para.text)

        # Extract tables
        for table in document.tables:
            for row in table.rows:
                row_text = '\t'.join(cell.text for cell in row.cells)
                full_text.append(row_text)

        return '\n\n'.join(full_text)  # 合并为一个字符串，段落之间空两行


    def extract_text_from_doc(doc_path):
        """Use antiword to extract text from a .doc file and return as a single string."""
        try:
            result = subprocess.run(['antiword', doc_path], stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            if result.returncode == 0:
                return result.stdout.decode('ISO-8859-1')
            else:
                print(f"Error processing {doc_path}: {result.stderr.decode()}")
                return ""
        except Exception as e:
            print(f"Exception while processing {doc_path}: {e}")
            return ""


    def dir_to_json(directory):
        """Convert all .doc or .docx files in the directory into a list of objects with filename and text."""
        all_files_content = []

        for filename in os.listdir(directory):
            if filename.endswith(('.doc', '.docx')):
                file_path = os.path.join(directory, filename)

                if filename.endswith('.docx'):
                    text = extract_text_from_docx(file_path)
                elif filename.endswith('.doc'):
                    text = extract_text_from_doc(file_path)

                all_files_content.append({
                    "text": text.strip()  # 去除首尾空白
                })

        return all_files_content  # 返回对象列表，每个对象包含一个文件的所有文本内容


    if __name__ == '__main__':
        import argparse
        parser = argparse.ArgumentParser(description='Convert all .doc or .docx files in a directory to a JSON file, each file as one field in an array.')
        parser.add_argument('directory', type=str, help='The path to the directory containing .doc or .docx files')
        parser.add_argument('output_json', type=str, help='The path where the output JSON file will be saved')

        args = parser.parse_args()

        combined_content = dir_to_json(args.directory)

        with open(args.output_json, 'w', encoding='utf-8') as json_file:
            json.dump(combined_content, json_file, ensure_ascii=False, indent=4)    
else:
        
    import os
    import json
    from docx import Document
    import subprocess

    def extract_pages_from_docx(docx_path, lines_per_page=40):
        """从 .docx 文件中提取页面内容，模拟分页"""
        document = Document(docx_path)
        pages = []
        current_page = []

        for para in document.paragraphs:
            # 检查是否是分页符
            if para._element.xml.find('w:br') >= 0 and 'type="page"' in para._element.xml:
                if current_page:
                    pages.append("\n".join(current_page).strip())
                    current_page = []
            else:
                if para.text.strip():
                    current_page.append(para.text.strip())

            # 如果当前页达到最大行数，强制分页
            if len(current_page) >= lines_per_page:
                pages.append("\n".join(current_page).strip())
                current_page = []

        # 添加最后一页
        if current_page:
            pages.append("\n".join(current_page).strip())

        return [p for p in pages if p]  # 跳过空页


    def extract_pages_from_doc(doc_path, lines_per_page=40):
        """使用 antiword 提取 .doc 内容，并模拟分页"""
        try:
            result = subprocess.run(['antiword', doc_path], stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            if result.returncode == 0:
                raw_text = result.stdout.decode('ISO-8859-1')
                lines = raw_text.splitlines()
                pages = []

                for i in range(0, len(lines), lines_per_page):
                    page_lines = lines[i:i + lines_per_page]
                    page_text = "\n".join(line.strip() for line in page_lines if line.strip())
                    if page_text:
                        pages.append(page_text)

                return pages
            else:
                print(f"Error processing {doc_path}: {result.stderr.decode()}")
                return []
        except Exception as e:
            print(f"Exception while processing {doc_path}: {e}")
            return []


    def process_directory(directory):
        all_pages = []

        for filename in os.listdir(directory):
            file_path = os.path.join(directory, filename)

            if filename.endswith('.docx'):
                print(f"Processing DOCX: {file_path}")
                pages = extract_pages_from_docx(file_path)
                all_pages.extend({"text": page} for page in pages)

            elif filename.endswith('.doc'):
                print(f"Processing DOC: {file_path}")
                pages = extract_pages_from_doc(file_path)
                all_pages.extend({"text": page} for page in pages)

        return all_pages


    if __name__ == "__main__":
        import argparse

        parser = argparse.ArgumentParser(description='Convert each page of .doc or .docx files in a directory to a JSON file.')
        parser.add_argument('directory', type=str, help='Path to the directory containing .doc or .docx files')
        parser.add_argument('output_json', type=str, help='Path to save the output JSON file')

        args = parser.parse_args()

        pages_data = process_directory(args.directory)

        with open(args.output_json, 'w', encoding='utf-8') as f:
            json.dump(pages_data, f, ensure_ascii=False, indent=4)

        print(f"Total non-empty pages extracted: {len(pages_data)}")